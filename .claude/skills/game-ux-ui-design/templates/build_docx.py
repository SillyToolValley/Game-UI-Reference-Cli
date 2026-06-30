#!/usr/bin/env python3
"""Render a UX/UI 디자인 문서 (markdown) to an EDITABLE landscape Word doc (.docx).

For when meeting participants need to edit the design document themselves (Word / Google Docs),
not just read a PDF. Markdown stays the source; this is the editable share format.

    pip install python-docx
    python build_docx.py design.md                 # -> design.docx next to it
    python build_docx.py design.md --append references/hud.md ... --out out/design.docx

Wireframes are embedded as their PNG renders (Word can't show SVG) — run build_pdf.py /
your SVG render first so wireframes/*.png exist. Reference .jpg/.png embed directly.
A line-based markdown reader handles the design's headings / tables / images / lists /
blockquotes / fenced code; landscape 16:9 page; Korean-capable font.
"""
import argparse, os, re, sys
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ACCENT = RGBColor(0x1E, 0x3A, 0x8A)
HEAD_BG = "1E3A8A"
ZEBRA = "F4F7FC"

LATIN = "JetBrains Mono"      # Latin / numbers / code
KOREAN = "Malgun Gothic"      # Hangul (JetBrains Mono has no Hangul glyphs)

def _set_rfonts(rpr, latin, korean):
    rf = rpr.find(qn('w:rFonts'))
    if rf is None:
        rf = OxmlElement('w:rFonts'); rpr.append(rf)
    rf.set(qn('w:ascii'), latin); rf.set(qn('w:hAnsi'), latin)
    rf.set(qn('w:cs'), latin); rf.set(qn('w:eastAsia'), korean)

def set_doc_fonts(doc, size=12):
    st = doc.styles["Normal"]; st.font.size = Pt(size)
    _set_rfonts(st.element.get_or_add_rPr(), LATIN, KOREAN)

def run_fonts(r, latin=LATIN, korean=KOREAN):
    _set_rfonts(r._element.get_or_add_rPr(), latin, korean)

def shade(cell_or_row_tc, hexcolor):
    tcPr = cell_or_row_tc._tc.get_or_add_tcPr()
    sh = OxmlElement('w:shd'); sh.set(qn('w:val'), 'clear'); sh.set(qn('w:fill'), hexcolor)
    tcPr.append(sh)

def add_runs(p, text, base_size=None):
    # inline **bold**, `code`, [text](url)->text; strip leftover md emphasis
    text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)
    text = text.replace('~~', '')
    for seg in re.split(r'(\*\*[^*]+\*\*|`[^`]+`)', text):
        if not seg: continue
        if seg.startswith('**') and seg.endswith('**'):
            r = p.add_run(seg[2:-2]); r.bold = True
        elif seg.startswith('`') and seg.endswith('`'):
            r = p.add_run(seg[1:-1]); r.font.name = 'JetBrains Mono'; r.font.size = Pt((base_size or 12) - 1)
        else:
            p.add_run(seg)
        if base_size:
            for r in p.runs[-1:]: r.font.size = Pt(base_size)

def emit_table(doc, rows):
    rows = [r for r in rows if not re.match(r'^\s*\|?[\s:|-]+\|?\s*$', r)]  # drop --- separator
    grid = [[c.strip() for c in re.split(r'(?<!\\)\|', r.strip().strip('|'))] for r in rows]
    ncol = max(len(r) for r in grid)
    grid = [r + [''] * (ncol - len(r)) for r in grid]
    t = doc.add_table(rows=len(grid), cols=ncol); t.style = 'Table Grid'; t.autofit = True
    for ri, row in enumerate(grid):
        for ci, val in enumerate(row):
            cell = t.cell(ri, ci); cell.paragraphs[0].clear()
            add_runs(cell.paragraphs[0], val.replace('\\|', '|'), base_size=10)
            if ri == 0:
                shade(cell, HEAD_BG)
                for r in cell.paragraphs[0].runs:
                    r.bold = True; r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            elif ri % 2 == 0:
                shade(cell, ZEBRA)
    doc.add_paragraph()

def emit_image(doc, src_abs):
    if src_abs.lower().endswith('.svg'):
        png = src_abs[:-4] + '.png'
        if os.path.exists(png): src_abs = png
        else:
            doc.add_paragraph(f"[wireframe: {os.path.basename(src_abs)} — PNG 없음, build_pdf.py 먼저]").italic = True
            return
    if not os.path.exists(src_abs):
        doc.add_paragraph(f"[이미지 없음: {os.path.basename(src_abs)}]"); return
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    wire = src_abs.lower().endswith('.png')
    try:
        p.add_run().add_picture(src_abs, width=Inches(14.5 if wire else 6.2))
    except Exception as e:
        doc.add_paragraph(f"[이미지 삽입 실패: {os.path.basename(src_abs)} — {e}]")

def md_to_docx(doc, md_text, src_dir):
    lines = md_text.splitlines()
    i = 0; in_code = False; code_buf = []
    while i < len(lines):
        ln = lines[i]
        if ln.strip().startswith('```'):
            if in_code:
                p = doc.add_paragraph(); shade_para(p, "EEF1F6")
                r = p.add_run('\n'.join(code_buf)); r.font.name = 'Consolas'; r.font.size = Pt(8.5)
                code_buf = []; in_code = False
            else:
                in_code = True
            i += 1; continue
        if in_code:
            code_buf.append(ln); i += 1; continue
        m = re.match(r'^(#{1,4})\s+(.*)$', ln)
        if m:
            lvl = len(m.group(1)); txt = m.group(2)
            h = doc.add_heading(level=min(lvl, 4)); h.clear()
            add_runs(h, txt)
            hsize = {1: 22, 2: 17, 3: 14.5, 4: 12.5}.get(lvl, 12)
            for r in h.runs:
                r.font.color.rgb = ACCENT; r.font.size = Pt(hsize); run_fonts(r)
            i += 1; continue
        if ln.strip().startswith('|') and i + 1 < len(lines) and re.match(r'^\s*\|?[\s:|-]+\|?\s*$', lines[i+1]):
            tbl = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                tbl.append(lines[i]); i += 1
            emit_table(doc, tbl); continue
        im = re.match(r'^!\[[^\]]*\]\(([^)]+)\)\s*$', ln.strip())
        if im:
            emit_image(doc, os.path.normpath(os.path.join(src_dir, im.group(1)))); i += 1; continue
        if ln.strip().startswith('>'):
            p = doc.add_paragraph(); p.paragraph_format.left_indent = Inches(0.2); shade_para(p, "F7F9FD")
            add_runs(p, ln.strip().lstrip('>').strip(), base_size=9)
            for r in p.runs: r.italic = True
            i += 1; continue
        lm = re.match(r'^(\s*)[-*]\s+(.*)$', ln)
        if lm:
            p = doc.add_paragraph(style='List Bullet'); add_runs(p, lm.group(2)); i += 1; continue
        if re.match(r'^-{3,}\s*$', ln.strip()):
            i += 1; continue
        if ln.strip():
            p = doc.add_paragraph(); add_runs(p, ln.strip())
        i += 1

def shade_para(p, hexc):
    pPr = p._p.get_or_add_pPr(); sh = OxmlElement('w:shd')
    sh.set(qn('w:val'), 'clear'); sh.set(qn('w:fill'), hexc); pPr.append(sh)

def build(md_path, out_docx, appends=None):
    md_path = os.path.abspath(md_path); src_dir = os.path.dirname(md_path)
    doc = Document()
    s = doc.sections[0]; s.orientation = WD_ORIENT.LANDSCAPE
    s.page_width, s.page_height = Inches(16), Inches(9)
    s.left_margin = s.right_margin = Inches(0.5); s.top_margin = s.bottom_margin = Inches(0.5)
    set_doc_fonts(doc, size=12)
    md_to_docx(doc, open(md_path, encoding='utf-8').read(), src_dir)
    for ap in (appends or []):
        ap = os.path.abspath(ap)
        if not os.path.exists(ap): continue
        doc.add_page_break()
        md_to_docx(doc, open(ap, encoding='utf-8').read(), os.path.dirname(ap))
    doc.save(out_docx)
    print("DOCX:", out_docx, f"({os.path.getsize(out_docx)//1024} KB)")

if __name__ == "__main__":
    ap = argparse.ArgumentParser()
    ap.add_argument("md"); ap.add_argument("--out", default=None)
    ap.add_argument("--append", action="append", default=[])
    a = ap.parse_args()
    out = a.out or os.path.splitext(os.path.abspath(a.md))[0] + ".docx"
    os.makedirs(os.path.dirname(out) or ".", exist_ok=True)
    build(a.md, out, appends=a.append)
